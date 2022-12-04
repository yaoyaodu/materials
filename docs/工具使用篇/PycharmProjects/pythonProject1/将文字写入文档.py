from docx import Document
doc = Document()
#添加标题。标题的等级通过level参数指定
doc.add_heading('一级标题', level=1)
doc.add_heading('二级标题', level=2)
doc.add_heading('三级标题', level=3)
#添加段落
p2 = doc.add_paragraph('第二个段落')
#将新段落添加到已有段落之前
p1 = p2.insert_paragraph_before('第一个段落')

p3 = doc.add_paragraph('新段落')

#追加内容
p3.add_run('加粗').bold = True
p3.add_run('以及')
p3.add_run('斜体').italic = True

doc.save('D:/Friday.docx')