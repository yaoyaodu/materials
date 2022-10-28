# 读取Word所有内容
# from docx import Document
# 文件 = Document('c:/练习3.docx')
# for 段落 in 文件.paragraphs:
# print(段落.text)

# 读取一级或二级标题
# from docx import Document
# 文件 = Document('D:/Python_Test/NDA.docx')
# for 段落 in 文件.paragraphs:
# 	if 段落.style.name == 'Heading 1':
# 		print(段落.text)

#读取所有标题（使用正则表达式）
# from docx import Document
# import re
# 文件 = Document('D:/Python_Test/NDA.docx')
# for 段落 in 文件.paragraphs:
# 	if re.match("^Heading \d+$",段落.style.name):
# 		print(段落.text)


# from docx.enum.style import WD_STYLE_TYPE
# from docx import Document
# 文件 = Document('D:/Python_Test/NDA.docx')
# 样式 = 文件.styles
# for i in 样式:
# 	if i.type==WD_STYLE_TYPE.PARAGRAPH:
# 		print(i.name)

# #添加标题
# from docx import Document
# 文件 = Document('D:/Python_Test/NDA.docx')
# 文件.add_heading("我是新增的二级标题2", level=2)
# 文件.save('D:/Python_Test/NDA.docx')

# from docx import Document
# 文件 = Document('D:/Python_Test/NDA.docx')
# 文件.add_paragraph("我是正文1234567")
# 文件.save('D:/Python_Test/NDA.docx')

# from docx import Document
# 文件 = Document('D:/Python_Test/NDA.docx')
# 文件.add_page_break()
# 文件.save('D:/Python_Test/NDA.docx')


# from docx import Document
# 文件 = Document('D:/Python_Test/NDA.docx')
# a = 文件.add_paragraph('')
# a.add_run('本文档内容仅供参考。沐曦不对本文档所含信息的准确性或完整性作任何明示或暗示的陈述或保证，也不对本文档中存在的任何错误承担责任').bold = True
# a.add_run('普通')
# a.add_run('斜体').italic = True
# 文件.save('D:/Python_Test/NDA.docx')

# from docx import Document
# 文件 = Document('D:/Python_Test/zh1.docx')
# print(len(文件.paragraphs))
# 段落 = 文件.paragraphs[19]
# 段落.text = "直到20世纪70年代，硅谷都是这场革命的中心，因为它们结合了科学知识、制造技术和有远见的商业思维。后面的删除"
# 文件.save('D:/Python_Test/zh1.docx')

# from docx import Document
# 文件 = Document('D:/Python_Test/zh1.docx')
# 第二个段落 = 文件.paragraphs[1] # 获取第二个段落
# 第二个段落.insert_paragraph_before('这是添加的新的第个段落')# 在第二个段落处插入
# 文件.save('D:/Python_Test/zh1.docx')


from docx import Document
from docx.shared import Pt, RGBColor              # 字号，颜色
from docx.oxml.ns import qn                       # 中文字体

# 文件 = Document('D:/Python_Test/zh1.docx')
# for 段落 in 文件.paragraphs:
# 	for 块 in 段落.runs:
# 		块.font.bold = True # 加粗
# 		块.font.italic = True # 斜体
# 		块.font.underline = True # 下划线
# 		块.font.strike = True # 删除线
# 		块.font.shadow = True # 阴影
# 		块.font.size = Pt(12)
# 		块.font.color.rgb = RGBColor(255,0,0)# 颜色
# 		块.font.name = 'Arial' # 英文字体设置
# 		块._element.rPr.rFonts.set(qn('w:eastAsia'),'思源黑体 CN Normal')   # 设置中文字体
# 文件.save('D:/Python_Test/zh1.docx')

# from docx import Document
# from docx.oxml.ns import qn # 中文字体
# 文件 = Document('D:/Python_Test/zh1.docx')
# 文件.styles['Normal'].font.name = 'Arial'# 设置英文字体
# 文件.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '思源黑体 CN Normal')  #设置中文字体
# 文件.save('D:/Python_Test/zh1.docx')

# from docx import Document
# from docx.oxml.ns import qn # 中文字体
# 文件 = Document('D:/Python_Test/zh3.docx')
# 文件.styles['Heading 2'].font.name = 'Arial'# 设置英文字体
# 文件.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), '思源黑体 CN Normal')  #设置中文字体：微软雅黑、思源黑体 CN Normal
# 文件.save('D:/Python_Test/zh3.docx')

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
文件 = Document('D:/Python_Test/zh2.docx')
for 段落 in 文件.paragraphs:
	if 段落.style.name=='Normal':
		段落.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
文件.save('D:/Python_Test/zh2.docx')





