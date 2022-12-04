from docx import Document
# 文件 = Document() 可以理解为 Document就是一个类，这个操作也就是实例化的过程，生成对象为：文件
文件 = Document('D:\Python_Test\zh1.docx')
# print(文件.paragraphs)
print(文件.paragraphs[0:2])
for i in 文件.paragraphs[0:2]:
    print(i)

# for 段落 in 文件.paragraphs:
#     print(段落)
# 文件.paragraphs#返回文档中每个段落集合，是一个列表，可以通过索引获取

#    print(段落.text)
# print(文件.paragraphs[0:2])


